import openpyxl
import os

from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.core.files.storage import FileSystemStorage
from .models import Student
from .forms import UploadFileForm
from docx import Document


def upload_files(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            excel_file = request.FILES['excel_file']
            word_template = request.FILES['word_template']
            fs = FileSystemStorage()
            excel_filename = fs.save(excel_file.name, excel_file)
            word_filename = fs.save(word_template.name, word_template)
            excel_filepath = fs.path(excel_filename)
            word_filepath = fs.path(word_filename)
            process_files(excel_filepath, word_filepath)
            return redirect('results:student_list')
    else:
        form = UploadFileForm()
    return render(request, 'results/upload.html', {'form': form})

def process_files(excel_filepath, word_filepath):
    workbook = openpyxl.load_workbook(excel_filepath)
    sheet = workbook.active

    for row in sheet.iter_rows(min_row=2, values_only=True):
        student_id, name, result = row
        student = Student.objects.create(student_id=student_id, name=name, result=result)
        generate_result_sheet(student, word_filepath)

def generate_result_sheet(student, template_path):
    document = Document(template_path)

    for paragraph in document.paragraphs:
        # Debugging information
        print("Original Paragraph:", paragraph.text)
        try:
            paragraph.text = paragraph.text.replace('{student_id}', str(student.student_id))
            paragraph.text = paragraph.text.replace('{name}', str(student.name))
            paragraph.text = paragraph.text.replace('{result}', str(student.result))
        except Exception as e:
            print("Error in replacing text:", e)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # Debugging information
                print("Original Cell Text:", cell.text)
                try:
                    cell.text = cell.text.replace('{student_id}', str(student.student_id))
                    cell.text = cell.text.replace('{name}', str(student.name))
                    cell.text = cell.text.replace('{result}', str(student.result))
                except Exception as e:
                    print("Error in replacing cell text:", e)

    output_path = os.path.join('media', 'results', f'{student.student_id}_result.docx')
    if not os.path.exists(os.path.dirname(output_path)):
        os.makedirs(os.path.dirname(output_path))
    document.save(output_path)


# -----------------------------------------------------

def student_list(request):
    students = Student.objects.all()
    return render(request, 'results/student_list.html', {'students': students})


def student_result(request, student_id):
    students = Student.objects.filter(student_id=student_id)
    if students.exists():
        student = students.first()  # Get the first result
        result_path = os.path.join('media', 'results', f'{student.student_id}_result.docx')
        return render(request, 'results/student_result.html', {'student': student, 'result_path': result_path})
    else:
        # Handle the case where no student is found
        return HttpResponse("Student not found", status=404)


#--------------PDF Download -----------------

import os
import win32com.client
from django.conf import settings
from django.http import HttpResponse
from PyPDF2 import PdfMerger
from docx2pdf import convert
from win32com.client import Dispatch

def download_all_results(request):
    # Ensure COM is initialized
    win32com.client.pythoncom.CoInitialize()

    output_dir = os.path.join(settings.MEDIA_ROOT, 'results')
    docx_files = [os.path.join(output_dir, f) for f in os.listdir(output_dir) if f.endswith('.docx')]

    if not docx_files:
        return HttpResponse("No result DOCX files found.", status=404)

    pdf_files = []

    # Convert DOCX to PDF
    for docx_file in docx_files:
        pdf_file = docx_file.replace('.docx', '.pdf')
        try:
            convert(docx_file, pdf_file)
            pdf_files.append(pdf_file)
        except Exception as e:
            return HttpResponse(f"Error converting {docx_file} to PDF: {e}", status=500)

    combined_pdf_path = os.path.join(output_dir, 'combined_results.pdf')

    # Merge all PDFs
    merger = PdfMerger()
    for pdf in pdf_files:
        merger.append(pdf)

    merger.write(combined_pdf_path)
    merger.close()

    # Serve the combined PDF
    with open(combined_pdf_path, 'rb') as pdf_file:
        response = HttpResponse(pdf_file.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="all_student_results.pdf"'
        return response
